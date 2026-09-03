# -*- coding: utf-8 -*-
"""
Docx 原始全量数据智能切片工具
将 9.3 版 3015 个段落拆解为两大候选集：
1. raw_skills_dump.json (包含 GitHub/Gitee/Skill名称/命令的候选块)
2. raw_styles_dump.json (包含 生图提示词/画幅/模型/动态视频提示词的候选块)
3. raw_tools_dump.json (包含 在线工具/设计工具库候选块)
"""

import os
import json
import re
import docx

DOCX_PATH = "网络热门风格➕Skill开源提示词(9.3).docx"

def is_skill_candidate(text):
    t = text.lower()
    if any(k in t for k in ["github.com", "gitee.com", "npx skills", "skills add", "git clone"]):
        return True
    if re.search(r'\bskill\b|【skill】|\[skill\]|开源skill|这个skill', t):
        return True
    return False

def is_tool_candidate(text):
    t = text.lower()
    if any(k in t for k in ["工具箱", "在线工具", "css 阴影", "调色板", "配色", "渐变", "huotu333", "lingmao88"]):
        if not is_skill_candidate(text):
            return True
    return False

def run_extraction():
    doc = docx.Document(DOCX_PATH)
    print(f"Total paragraphs in docx: {len(doc.paragraphs)}")

    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            paragraphs.append({
                "idx": i,
                "style": p.style.name,
                "text": t
            })

    print(f"Non-empty paragraphs: {len(paragraphs)}")

    # 1. First pass: group by headers and logical blocks
    blocks = []
    curr = []
    for p in paragraphs:
        t = p["text"]
        # If paragraph is a standalone number like "1", "2", "3" or Heading or section title
        if p["style"] != "Normal" or re.match(r'^\d+$', t) or (len(t) < 40 and any(k in t for k in ["风格", "提示词", "Skill", "skill", "工具", "合集", "Part", "一、", "二、", "三、", "四、"])):
            if curr:
                blocks.append(curr)
                curr = []
        curr.append(p)
    if curr:
        blocks.append(curr)

    print(f"Aggregated into {len(blocks)} coarse blocks.")

    skills_raw = []
    tools_raw = []
    styles_raw = []

    for b in blocks:
        combined_text = "\n".join(p["text"] for p in b)
        start_idx = b[0]["idx"]
        end_idx = b[-1]["idx"]

        if is_skill_candidate(combined_text):
            skills_raw.append({
                "start_idx": start_idx,
                "end_idx": end_idx,
                "paragraphs": [p["text"] for p in b],
                "combined": combined_text
            })
        elif is_tool_candidate(combined_text):
            tools_raw.append({
                "start_idx": start_idx,
                "end_idx": end_idx,
                "paragraphs": [p["text"] for p in b],
                "combined": combined_text
            })
        else:
            # Check if it's a style prompt
            if len(combined_text) > 35 and not combined_text.startswith("《提示词禁止商用"):
                styles_raw.append({
                    "start_idx": start_idx,
                    "end_idx": end_idx,
                    "paragraphs": [p["text"] for p in b],
                    "combined": combined_text
                })

    print(f"Extracted candidate Skills: {len(skills_raw)}")
    print(f"Extracted candidate Tools: {len(tools_raw)}")
    print(f"Extracted candidate Styles: {len(styles_raw)}")

    with open("scripts/raw_skills_dump.json", "w", encoding="utf-8") as f:
        json.dump(skills_raw, f, ensure_ascii=False, indent=2)

    with open("scripts/raw_tools_dump.json", "w", encoding="utf-8") as f:
        json.dump(tools_raw, f, ensure_ascii=False, indent=2)

    with open("scripts/raw_styles_dump.json", "w", encoding="utf-8") as f:
        json.dump(styles_raw, f, ensure_ascii=False, indent=2)

    print("Dumps written successfully to scripts/raw_*_dump.json")

if __name__ == "__main__":
    run_extraction()
