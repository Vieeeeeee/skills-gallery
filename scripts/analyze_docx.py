import docx

doc = docx.Document("网络热门风格➕Skill开源提示词(3).docx")
lines = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        lines.append(f"[{i}] {t}")

with open("scripts/dump_docx.txt", "w", encoding="utf-8") as f:
    f.write("
".join(lines))

print(f"Dumped {len(lines)} non-empty paragraphs from {len(doc.paragraphs)} total paragraphs")