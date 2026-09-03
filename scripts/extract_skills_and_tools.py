# -*- coding: utf-8 -*-
import docx
import re
import json

doc = docx.Document("网络热门风格➕Skill开源提示词(9.3).docx")

all_lines = []
for p_idx in range(2045, len(doc.paragraphs)):
    p = doc.paragraphs[p_idx]
    raw = p.text.replace("\r", "\n").replace("\x0b", "\n")
    for l in raw.split("\n"):
        t = l.strip()
        all_lines.append((p_idx, t))

# Collect all skill and tool blocks
# A skill is recognized by github URL, gitee URL, or explicit skill name marker
skills_dict = {}

# 1. First pass: find GitHub / Gitee repos
for i, (p_idx, line) in enumerate(all_lines):
    gh_match = re.search(r'https?://(?:www\.)?github\.com/([a-zA-Z0-9_\-\.]+)/([a-zA-Z0-9_\-\.]+)', line)
    if gh_match:
        owner = gh_match.group(1)
        repo = gh_match.group(2).rstrip('.git').rstrip('/').rstrip('）').rstrip(')')
        # Clean common typos
        if repo.endswith('.gitcd'):
            repo = repo[:-6]
        if repo == 'gc-minimal-':
            repo = 'gc-minimal-zine-poster'
        if repo == 'desktop-pe':
            repo = 'desktop-pet'
        if repo == 'html-anythin':
            repo = 'html-anything'
        if repo == 'Starryear-Abstract-Quarte':
            repo = 'Starryear-Abstract-Quartet'
        if repo == 'Legal-':
            repo = 'legal-skills'
            
        full_key = f"{owner}/{repo}".lower()
        
        # Context lines around this mention
        before = [all_lines[j][1] for j in range(max(0, i-4), i) if all_lines[j][1]]
        after = [all_lines[j][1] for j in range(i+1, min(len(all_lines), i+8)) if all_lines[j][1]]
        
        skills_dict.setdefault(full_key, {
            "owner": owner,
            "repo": repo,
            "url": f"https://github.com/{owner}/{repo}",
            "mentions": [],
            "raw_contexts": []
        })
        skills_dict[full_key]["mentions"].append(line)
        skills_dict[full_key]["raw_contexts"].append({"before": before, "after": after})

print(f"Total distinct GitHub repos found: {len(skills_dict)}")
for k, v in list(skills_dict.items())[:20]:
    print(f"  {k} -> mentions: {len(v['mentions'])}")
